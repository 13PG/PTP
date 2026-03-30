from __future__ import annotations

import io
import logging
import re
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import fitz
from docx import Document as DocxDocument
from PIL import Image, UnidentifiedImageError

from app.services.ocr_support import extract_page_text_with_ocr, is_ocr_available

logger = logging.getLogger(__name__)
fitz.TOOLS.mupdf_display_errors(False)
fitz.TOOLS.mupdf_display_warnings(False)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}


@dataclass
class ExtractedImage:
    path: Path
    origin: str
    caption: str = ""
    width: int = 0
    height: int = 0
    score: float = 0.0


@dataclass
class TextChunk:
    text: str
    location: str


@dataclass
class ParsedDocument:
    source_path: Path
    file_type: str
    filename: str
    title: str
    authors: str
    abstract: str
    full_text: str
    chunks: list[TextChunk] = field(default_factory=list)
    images: list[ExtractedImage] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)


class DocumentParseError(RuntimeError):
    """Raised when a file cannot be parsed."""


def parse_document(source_path: Path, asset_dir: Path) -> ParsedDocument:
    extension = source_path.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise DocumentParseError(f"暂不支持的文件类型: {source_path.suffix}")

    asset_dir.mkdir(parents=True, exist_ok=True)

    if extension == ".pdf":
        return _parse_pdf(source_path, asset_dir)
    if extension == ".docx":
        return _parse_docx(source_path, asset_dir)

    converted = _convert_doc_to_docx(source_path, asset_dir)
    if converted is None:
        raise DocumentParseError(
            "检测到 .doc 文件，但当前环境无法自动转换。请安装 Microsoft Word 与 pywin32，或者先另存为 .docx。"
        )
    return _parse_docx(converted, asset_dir)


def _parse_pdf(source_path: Path, asset_dir: Path) -> ParsedDocument:
    fitz.TOOLS.reset_mupdf_warnings()
    try:
        document = fitz.open(source_path)
    except Exception as exc:
        raise DocumentParseError(f"无法打开 PDF 文件: {source_path.name}") from exc

    try:
        metadata = document.metadata or {}
        pages_text: list[str] = []
        chunks: list[TextChunk] = []
        image_candidates: list[ExtractedImage] = []
        seen_xrefs: set[int] = set()

        for page_index, page in enumerate(document):
            page_text = _extract_pdf_page_text(page)
            figure_caption = _guess_figure_caption(page_text)
            has_figure_caption = bool(figure_caption)
            if page_text:
                pages_text.append(page_text)
                chunks.extend(_split_into_chunks(page_text, f"第 {page_index + 1} 页"))

            for image_index, image_info in enumerate(page.get_images(full=True), start=1):
                xref = image_info[0]
                if xref in seen_xrefs:
                    continue
                seen_xrefs.add(xref)

                try:
                    base_image = document.extract_image(xref)
                except Exception:
                    logger.exception("提取 PDF 图片失败: %s xref=%s", source_path, xref)
                    continue

                image_path, width, height = _save_image_bytes(
                    base_image.get("image"),
                    base_image.get("ext", "png"),
                    asset_dir / f"{source_path.stem}_page{page_index + 1}_{image_index}",
                )
                if image_path is None:
                    continue
                if width < 180 or height < 180:
                    image_path.unlink(missing_ok=True)
                    continue
                relative_area = _estimate_pdf_image_area(page, xref)
                score = _score_pdf_image(width, height, relative_area, has_figure_caption, page_index)
                if score < 2.6:
                    image_path.unlink(missing_ok=True)
                    continue

                image_candidates.append(
                    ExtractedImage(
                        path=image_path,
                        origin=f"第 {page_index + 1} 页",
                        caption=figure_caption or "论文实验配图",
                        width=width,
                        height=height,
                        score=score,
                    )
                )

        full_text = "\n\n".join(pages_text)
        first_page_text = pages_text[0] if pages_text else ""
        title = _first_non_empty(metadata.get("title"), _guess_title(first_page_text), source_path.stem)
        authors = _first_non_empty(metadata.get("author"), _guess_authors(first_page_text, title), "未识别作者")
        abstract = _extract_abstract("\n".join(pages_text[:3]))
        images = _pick_top_images(image_candidates)

        return ParsedDocument(
            source_path=source_path,
            file_type="pdf",
            filename=source_path.name,
            title=title,
            authors=authors,
            abstract=abstract,
            full_text=full_text,
            chunks=chunks,
            images=images,
            metadata=metadata,
        )
    finally:
        document.close()


def _parse_docx(source_path: Path, asset_dir: Path) -> ParsedDocument:
    try:
        doc = DocxDocument(str(source_path))
    except Exception as exc:
        raise DocumentParseError(f"无法打开 Word 文件: {source_path.name}") from exc

    paragraphs = [_clean_text(paragraph.text) for paragraph in doc.paragraphs]
    paragraphs = [paragraph for paragraph in paragraphs if paragraph]
    full_text = "\n\n".join(paragraphs)
    chunks = _split_paragraphs(paragraphs)
    core = doc.core_properties

    title = _first_non_empty(core.title, _guess_title(full_text), source_path.stem)
    authors = _first_non_empty(core.author, _guess_authors(full_text, title), "未识别作者")
    abstract = _extract_abstract(full_text)
    images = _extract_docx_images(source_path, asset_dir)

    metadata = {
        "title": core.title,
        "author": core.author,
        "subject": core.subject,
        "created": core.created.isoformat() if core.created else "",
        "modified": core.modified.isoformat() if core.modified else "",
    }

    return ParsedDocument(
        source_path=source_path,
        file_type="docx",
        filename=source_path.name,
        title=title,
        authors=authors,
        abstract=abstract,
        full_text=full_text,
        chunks=chunks,
        images=images,
        metadata=metadata,
    )


def _convert_doc_to_docx(source_path: Path, asset_dir: Path) -> Path | None:
    try:
        import win32com.client  # type: ignore
    except Exception:
        return None

    target_path = asset_dir / f"{source_path.stem}_converted.docx"
    word = None
    doc = None

    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(source_path.resolve()))
        doc.SaveAs(str(target_path.resolve()), FileFormat=16)
        return target_path
    except Exception:
        logger.exception("DOC 转 DOCX 失败: %s", source_path)
        return None
    finally:
        if doc is not None:
            doc.Close(False)
        if word is not None:
            word.Quit()


def _extract_docx_images(source_path: Path, asset_dir: Path) -> list[ExtractedImage]:
    image_candidates: list[ExtractedImage] = []

    try:
        with zipfile.ZipFile(source_path) as archive:
            media_files = [
                name
                for name in archive.namelist()
                if name.startswith("word/media/") and not name.endswith("/")
            ]
            for index, member in enumerate(media_files[:4], start=1):
                raw = archive.read(member)
                suffix = Path(member).suffix or ".png"
                image_path, width, height = _save_image_bytes(
                    raw,
                    suffix.lstrip("."),
                    asset_dir / f"{source_path.stem}_image{index}",
                )
                if image_path is None:
                    continue
                if width < 180 or height < 180:
                    image_path.unlink(missing_ok=True)
                    continue
                score = _score_docx_image(width, height, index)
                if score < 2.1:
                    image_path.unlink(missing_ok=True)
                    continue
                image_candidates.append(
                    ExtractedImage(
                        path=image_path,
                        origin=f"Word 图片 {index}",
                        caption="文档实验配图",
                        width=width,
                        height=height,
                        score=score,
                    )
                )
    except zipfile.BadZipFile as exc:
        raise DocumentParseError(f"Word 文件结构异常: {source_path.name}") from exc

    return _pick_top_images(image_candidates)


def _save_image_bytes(raw: bytes | None, ext: str, target_base: Path) -> tuple[Path | None, int, int]:
    if not raw:
        return None, 0, 0

    ext = ext.lower().strip(".")
    target_path = target_base.parent / f"{target_base.name}.png"

    try:
        with Image.open(io.BytesIO(raw)) as image:
            image.load()
            if image.mode in {"CMYK", "P"}:
                image = image.convert("RGB")
            if image.mode == "RGBA":
                background = Image.new("RGB", image.size, (255, 255, 255))
                background.paste(image, mask=image.split()[-1])
                image = background
            elif image.mode != "RGB":
                image = image.convert("RGB")
            image.save(target_path, format="PNG")
            return target_path, image.width, image.height
    except UnidentifiedImageError:
        logger.warning("无法识别图片格式 ext=%s", ext)
        return None, 0, 0
    except Exception:
        logger.exception("保存图片失败: %s", target_base)
        return None, 0, 0


def _split_paragraphs(paragraphs: list[str]) -> list[TextChunk]:
    chunks: list[TextChunk] = []
    for index, paragraph in enumerate(paragraphs, start=1):
        if len(paragraph) < 40:
            continue
        chunks.append(TextChunk(text=paragraph, location=f"段落 {index}"))
    return chunks


def _split_into_chunks(text: str, location_prefix: str) -> list[TextChunk]:
    segments = re.split(r"\n{2,}", text)
    chunks: list[TextChunk] = []
    for index, segment in enumerate(segments, start=1):
        cleaned = _clean_text(segment)
        if len(cleaned) < 40:
            continue
        chunks.append(TextChunk(text=cleaned, location=f"{location_prefix} / 片段 {index}"))
    return chunks


def _extract_abstract(text: str) -> str:
    patterns = [
        r"(abstract[:\s].{60,1600}?)(?:\n\s*(?:keywords?|introduction|1\.)\b)",
        r"(摘要[:：\s].{30,1200}?)(?:\n\s*(?:关键词|关键字|引言|1[\.、]))",
    ]
    for pattern in patterns:
        match = re.search(pattern, text or "", flags=re.IGNORECASE | re.DOTALL)
        if match:
            return _trim_text(_clean_text(match.group(1)), 420)
    return _trim_text(_clean_text((text or "")[:600]), 420)


def _guess_title(text: str) -> str:
    lines = [line for line in (_clean_text(line) for line in text.splitlines()) if line]
    for line in lines[:12]:
        lower = line.lower()
        if lower.startswith(("abstract", "摘要", "keywords", "关键词")):
            continue
        if "@" in line:
            continue
        if 12 <= len(line) <= 220:
            return line
    return ""


def _guess_authors(text: str, title: str) -> str:
    lines = [line for line in (_clean_text(line) for line in text.splitlines()) if line]
    for index, line in enumerate(lines[:15]):
        if title and line == title:
            window = lines[index + 1 : index + 5]
            for candidate in window:
                lower = candidate.lower()
                if lower.startswith(("abstract", "摘要", "keywords", "关键词")):
                    break
                if len(candidate) > 120:
                    continue
                if any(marker in candidate for marker in [",", ";", " and ", "、"]) or "@" in candidate:
                    return candidate.replace(";", ",")
                if 4 <= len(candidate.split()) <= 12:
                    return candidate

    for line in lines[:10]:
        if any(marker in line for marker in [",", " and ", "、"]):
            return line
    return ""


def _guess_figure_caption(page_text: str) -> str:
    patterns = [
        r"(Figure\s+\d+[^\n]{0,120})",
        r"(Fig\.\s*\d+[^\n]{0,120})",
        r"(图\s*\d+[^\n]{0,120})",
    ]
    for pattern in patterns:
        match = re.search(pattern, page_text or "", flags=re.IGNORECASE)
        if match:
            return _trim_text(_clean_text(match.group(1)), 90)
    return ""


def _clean_text(text: str | None) -> str:
    if not text:
        return ""
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s+\n", "\n\n", text)
    return text.strip()


def _trim_text(text: str, limit: int) -> str:
    if len(text) <= limit:
        return text
    return text[: limit - 1].rstrip() + "…"


def _first_non_empty(*values: Any) -> str:
    for value in values:
        if isinstance(value, str) and value.strip():
            return value.strip()
    return ""


def _estimate_pdf_image_area(page: fitz.Page, xref: int) -> float:
    try:
        rects = page.get_image_rects(xref)
    except Exception:
        return 0.0

    if not rects:
        return 0.0

    image_area = max(rect.width * rect.height for rect in rects)
    page_area = max(page.rect.width * page.rect.height, 1.0)
    return image_area / page_area


def _score_pdf_image(
    width: int,
    height: int,
    relative_area: float,
    has_figure_caption: bool,
    page_index: int,
) -> float:
    pixel_area = width * height
    aspect_ratio = width / max(height, 1)
    score = min(pixel_area / 250000.0, 6.0)
    score += min(relative_area * 14.0, 4.5)

    if has_figure_caption:
        score += 4.0
    if page_index > 0:
        score += 0.6
    if 0.45 <= aspect_ratio <= 2.8:
        score += 1.2
    if 0.85 <= aspect_ratio <= 1.15 and not has_figure_caption:
        score -= 2.2
    if page_index == 0 and not has_figure_caption:
        score -= 2.5
    if relative_area < 0.03 and not has_figure_caption:
        score -= 2.0

    return score


def _score_docx_image(width: int, height: int, index: int) -> float:
    pixel_area = width * height
    aspect_ratio = width / max(height, 1)
    score = min(pixel_area / 250000.0, 6.0)

    if 0.45 <= aspect_ratio <= 2.8:
        score += 1.0
    else:
        score -= 0.8

    if 0.85 <= aspect_ratio <= 1.15 and pixel_area < 300000:
        score -= 2.0
    if index == 1 and 0.85 <= aspect_ratio <= 1.15 and pixel_area < 450000:
        score -= 1.2

    return score


def _pick_top_images(candidates: list[ExtractedImage]) -> list[ExtractedImage]:
    unique_candidates: list[ExtractedImage] = []
    seen_paths: set[Path] = set()
    for item in sorted(candidates, key=lambda image: image.score, reverse=True):
        if item.path in seen_paths:
            item.path.unlink(missing_ok=True)
            continue
        seen_paths.add(item.path)
        unique_candidates.append(item)

    selected = unique_candidates[:4]
    discarded = unique_candidates[4:]
    for item in discarded:
        item.path.unlink(missing_ok=True)
    return selected


def _extract_pdf_page_text(page: fitz.Page) -> str:
    page_text = _clean_text(page.get_text("text"))
    if not _needs_ocr(page_text):
        return page_text

    ocr_text = _clean_text(extract_page_text_with_ocr(page)) if is_ocr_available() else ""
    if _is_better_text(ocr_text, page_text):
        return ocr_text
    return page_text


def _needs_ocr(text: str) -> bool:
    cleaned = _clean_text(text)
    if not cleaned:
        return True

    usable = len(re.findall(r"[A-Za-z0-9\u4e00-\u9fff]", cleaned))
    quality = usable / max(len(cleaned), 1)
    question_ratio = cleaned.count("?") / max(len(cleaned), 1)
    control_ratio = len(re.findall(r"[\x00-\x08\x0b-\x1f]", cleaned)) / max(len(cleaned), 1)

    return usable < 45 or quality < 0.42 or question_ratio > 0.08 or control_ratio > 0.03


def _is_better_text(candidate: str, baseline: str) -> bool:
    if not candidate:
        return False
    if not baseline:
        return True

    candidate_score = _text_quality_score(candidate)
    baseline_score = _text_quality_score(baseline)
    return candidate_score > baseline_score * 1.15 or (
        candidate_score >= baseline_score * 0.95 and len(candidate) > len(baseline) * 1.2
    )


def _text_quality_score(text: str) -> float:
    cleaned = _clean_text(text)
    if not cleaned:
        return 0.0
    usable = len(re.findall(r"[A-Za-z0-9\u4e00-\u9fff]", cleaned))
    penalty = cleaned.count("?") * 2 + len(re.findall(r"[\x00-\x08\x0b-\x1f]", cleaned)) * 3
    return usable - penalty
